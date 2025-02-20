import os
from docx import Document
from docx.shared import Inches
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 将图片文件夹保存到word文档
def save_images_to_docx(input_directory):
    # 创建一个新的Word文档
    doc = Document()
    
    # 获取文件夹名称作为文档名称
    folder_name = os.path.basename(input_directory)
    
    # 获取文件夹中的所有图片文件，按照 image_x.jpg 的序列号排序
    image_files = [f for f in os.listdir(input_directory) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')) and f.startswith('image_')]
    image_files.sort(key=lambda x: int(x.split('_')[1].split('.')[0]))
    
    # 遍历图片并添加到文档
    for index, image_file in enumerate(image_files, 1):
        # 构建完整的图片路径
        image_path = os.path.join(input_directory, image_file)
        
        try:
            # 打开图片并获取尺寸
            with Image.open(image_path) as img:
                # 添加图片到文档
                doc.add_picture(image_path, width=Inches(6.0))
                
                # 在每张图片后添加分页符
                doc.add_page_break()
                
                # 在页脚中间添加页码
                section = doc.sections[-1]
                footer = section.footer
                footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                footer_paragraph.text = str(index)
        
        except Exception as e:
            print(f"处理图片 {image_file} 时出错: {e}")
    
    # 保存文档
    output_file = os.path.join(os.path.dirname(input_directory), f"{folder_name}.docx")
    doc.save(output_file)
    print(f"文档已保存：{output_file}")

# 提示用户输入图片文件夹路径
input_directory = input('请输入包含图片的文件夹路径：')
save_images_to_docx(input_directory)
